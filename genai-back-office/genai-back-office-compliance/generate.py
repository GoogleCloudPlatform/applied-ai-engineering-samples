from vertex_search_retrieval import vertex_search_response
from typing import Any, Generator
from google.cloud import bigquery

from docx import Document
import pandas as pd
import math

# import google.auth
# credentials, project = google.auth.default()




def read_docx(path): 
    """
    Reads a docx from the specified path and returns the text as list. 
    """
    doc_text = list() 
    document = Document(path)
    document_paragraphs = (document.paragraphs)

    for p in document_paragraphs: 
        doc_text.append(p.text)
    
    return doc_text 



def get_questions(doc_path):
    """
    Parses the CDP document and retrieves the questions. 
    Splits it at each instance of 'Q' in the doc. 
    """
    doc_text = read_docx(doc_path)
    doc_string = ''.join(str(e) for e in doc_text)
    questions = doc_string.split("(Q)")
    questions = questions[1::]
    return questions 



def make_docx(MODEL, filename): 
    """
    Reads the CDP document, parses it, feeds the questions to LLM and puts the answers back in a File. 
    The file is saved under Files/result.docx 
    """

    # Fetch questions from doc 
    questions = get_questions(filename)

    # Instantiate doc 
    doc = Document(filename)

    for question in questions: 

        model_response = vertex_search_response(MODEL, search_query=question)

        # Replace >> in doc with the model answer 
        found = False 
        for p in doc.paragraphs:
            if not found: 
                if '>>' in p.text:
                    text = p.text.replace('>>', model_response)
                    p.text = text 

                    print (p.text)
                    found = True 
                    #break
            

    doc.save('Files/result.docx')

    return ('Files/result.docx')

