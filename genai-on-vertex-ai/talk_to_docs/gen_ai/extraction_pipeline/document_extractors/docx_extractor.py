# Copyright 2024 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     https://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
"""
Provides the DocxExtractor class for extracting textual data from Docx (.docx) files and organizing 
the extracted data into separate files for structured document processing.
"""

import datetime
import json
import os
import re
import xml.etree.ElementTree as ET
import zipfile

import docx
from gen_ai.extraction_pipeline.document_extractors.base_extractor import BaseExtractor


# Functions to convert xml to txt
nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def qn(tag: str) -> str:
    """Stands for "qualified name".

    This utility function converts a familiar namespace-prefixed tag name like
    "w:p" into a Clark-notation qualified tag name for lxml. For example,
    `qn("w:p")` returns
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p".
    Source: https://github.com/python-openxml/python-docx/
    """
    prefix, tagroot = tag.split(":")
    uri = nsmap[prefix]
    return "{%s}%s" % (uri, tagroot)


def convert_xml_to_text(xml: str) -> str:
    """Converts XML with specific 'w:t', 'w:tab', 'w:br', 'w:cr', and 'w:p' tags into text.

    Uses a namespace map for tag resolution.

    Args:
        xml (str): The XML string to be converted.

    Returns:
        str: The extracted text representation of the XML content.
    """

    root = ET.fromstring(xml)
    text = ""

    # Define a mapping for easy tag handling
    tag_map = {
        qn("w:t"): lambda item: item.text if item.text is not None else "",
        qn("w:tab"): lambda _: "\t",
        qn("w:br"): lambda _: "\n",
        qn("w:cr"): lambda _: "\n",
        qn("w:p"): lambda _: "\n\n",
    }

    for item in root.iter():
        handler = tag_map.get(item.tag)
        if handler:
            text += handler(item)

    return text


def _convert_date(date_str: str) -> str | None:
    """Converts a date string in the format "Month Day, Year" into the format "MM/DD/YYYY".

    Args:
        date_str (str): The date string to convert.

    Returns:
        str: The reformatted date string in "MM/DD/YYYY" format if conversion is
        successful.
        None: If the provided date_str cannot be parsed into the expected
        format.
    """
    formats = ["%B %d, %Y", "%m/%d/%Y"]
    for fmt in formats:
        try:
            return datetime.datetime.strptime(date_str, fmt).strftime("%m/%d/%Y")
        except ValueError:
            continue  # If it doesn't match this format, try the next one
    return None  # If no format matches, return None


class DocxExtractor(BaseExtractor):
    """Extractor class of textual data from Doc (.docx) files and chunks sections into separate files.

    This class inherits from the `BaseExtractor` and provides specialized
    functionality for extracting text content from .docx documents.

    Args:
        filepath (str): The path to the .docx file.
        config_file_parameters (dict[str, str]): Configuration settings for the
          extraction process.

    Attributes:
        filepath (str): Stores the path to the input file.
        config_file_parameters (dict[str, str]): Stores the configuration
          parameters.
        docx_extraction (str): Configuration parameter fot the extraction
          method. Defaults to 'default'.
        docx_chunking (str):  Configuration parameter fot the chunking method.
          Defaults to 'default'.
    """

    def __init__(self, filepath: str, config_file_parameters: dict[str, str]):
        super().__init__(filepath, config_file_parameters)
        self.docx_extraction = config_file_parameters.get(
            "docx_extraction", "default"
        )
        self.docx_chunking = config_file_parameters.get(
            "docx_chunking", "default"
        )
        self.chunk_level = config_file_parameters.get(
            "docx_chunk_level", 1
        )
        self.document = None
        self.raw_text = None

    def create_filepath(
        self,
        metadata: dict[str, str | None],
        section_name: str,
        output_dir: str,
    ) -> str:
        """Creates a file-system-friendly path for saving a document section.

        * Removes extension of the filename
        * Removes leading digits, whitespace, and hyphens from the section name.
        * Condenses multiple underscores into single underscores.

        Args:
            metadata (dict[str, str | None]): A dictionary containing document
              metadata, including a 'filename' key.
            section_name (str): The name of the section being saved.
            output_dir (str): The directory where the generated file should be
              saved.

        Returns:
            str: A filepath constructed from the provided information.
        """

        filename = os.path.splitext(metadata["filename"])[0]
        filename += f"-{section_name.lower()}"
        filename = re.sub(r"[^\w.-]", "_", filename)
        filename = re.sub(r"__+", "_", filename).rstrip("_")
        filepath = os.path.join(output_dir, filename)
        return filepath

    def create_files(
        self,
        document_chunks: dict[tuple[int, str], str],
        metadata: dict[str, str | None],
        output_dir: str,
    ) -> bool:
        """Saves document sections and associated metadata to individual files.

        The function iterates over a dictionary of document chunks, generates
        filepaths, and writes both the text content and a corresponding metadata
        JSON file output directory.

        Args:
            document_chunks (dict[tuple[int, str], str]): A dictionary where
              keys are tuples of (section_id, section_title) and values are the
              corresponding text content.
            metadata (dict[str, str | None]):  Metadata associated with the
              overall document.
            output_dir (str):  The target directory for saving the output files.

        Returns:
            bool: True to indicate successful file creation.
        """
        for (_, section_name), context in document_chunks.items():
            filepath = self.create_filepath(metadata, section_name, output_dir)

            if context.strip():
                with open(filepath + ".txt", "w", encoding="utf-8") as f:
                    f.write(context)
                temp_metadata = metadata.copy()
                temp_metadata.pop("filename")
                temp_metadata["section_name"] = section_name.lower()

                with open(filepath + "_metadata.json", "w", encoding="utf-8") as f:
                    json.dump(temp_metadata, f)
        return True

    def process(self, output_dir: str) -> bool:
        """
        Main function that controls the processing of a .docx document, including extraction,
        metadata creation, chunking, and file saving.

        This function coordinates the key steps for processing a .docx document.
        It handles document extraction, metadata generation, applies  document
        chunking strategies, and saves the resulting chunks and metadata to
        files.

        Args:
            output_dir (str): The directory where the processed files should be
              saved.

        Returns:
            bool: True if the document processing was successful, False
            otherwise.
        """
        extractor = DefaultDocxExtractor(self.filepath)
        self.document = extractor.extract_document()
        self.raw_text = extractor.extract_text()

        metadata_creator = CustomKcDocxMetadataCreator(
            self.filepath, self.document, self.raw_text
        )
        metadata = metadata_creator.create_metadata()
        if not metadata:
            return False

        document_chunker = DefaultDocxChunker(self.document, self.raw_text, self.chunk_level)
        document_chunks = document_chunker.chunk_the_document()
        additional_kc_chunks = (
            CustomKcDocxChunker(self.document, self.raw_text).chunk_the_document()
            if self.docx_chunking == "custom"
            else None
        )
        if additional_kc_chunks:
            document_chunks.update(additional_kc_chunks)

        if not document_chunks:
            return False

        if self.create_files(document_chunks, metadata, output_dir):
            return True
        return False


class DefaultDocxExtractor:
    """Default extractor class that provides methods for extracting content from .docx files.

    Args:
        filepath (str): The path to the .docx file.

    Attributes:
        filepath (str): The path to the .docx file.
    """

    def __init__(self, filepath: str):
        self.filepath = filepath

    def extract_document(self) -> docx.Document:
        """Extracts the Document object from the .docx file.

        Returns:
            docx.Document: The extracted Document object.
        """
        document = docx.Document(self.filepath)
        return document

    def extract_text(self) -> str:
        """Extracts the raw text content from the .docx file.

        Returns:
            str: The extracted text.
        """
        zipf = zipfile.ZipFile(self.filepath)
        text = convert_xml_to_text(zipf.read("word/document.xml"))
        return text


class DefaultDocxMetadataCreator:
    """Default class for creating metadata from .docx files.

    Provides a basic metadata structure including the filename.

    Args:
        filepath (str): The path to the .docx file.

    Attributes:
        filename (str): The name of the .docx file (without the path).
    """

    def __init__(self, filepath: str):
        filename = os.path.basename(filepath)
        self.filename = filename

    def create_metadata(self) -> dict[str, str]:
        """A method to be implemented by subclasses.

        Generates a dictionary of metadata extracted from the .docx file.

        Returns:
            dict[str, str]: A dictionary containing metadata keys and their
            corresponding values.
        """
        metadata = {
            "filename": self.filename,
        }
        return metadata


class CustomKcDocxMetadataCreator(DefaultDocxMetadataCreator):
    """Metadata creator from .docx class customly created for Custom-KC use case.

    Provides a basic metadata structure including the filename, policy name,
    title, etc.

    Args:
        filepath (str): The path to the .docx file.

    Attributes:
        filename (str): The name of the .docx file.
        document (docx.Document): The parsed Document object.
        raw_text (str): The raw text content extracted from the .docx file.
    """

    def __init__(self, filepath: str, document: docx.Document, raw_text: str):
        super().__init__(filepath)
        self.document = document
        self.raw_text = raw_text

    def _extract_policy_number(self) -> str | None:
        """Extracts the "Policy Number" from the raw text of the document.

        Searches for a pattern like "Policy Number: ...".

        Returns:
            str: The policy number if found.
            None: If the "Policy Number" pattern is not found in the raw text.
        """
        match = re.search(r"Policy Number:\s*([\w\d\.-]+)\n", self.raw_text)

        if match:
            policy_number = match.group(1).strip()
            return policy_number.lower()
        else:
            return None

    def _extract_effective_date(self) -> str | None:
        """Extracts the "Effective Date" from the raw text of the document.

        Searches for a pattern like "Effective Date: ...".

        Returns:
            str: The formatted effective date if found.
            None: If the "Effective Date" pattern is not found in the raw text.
        """
        match = re.search(r"Effective Date:\s*(.*?)\n", self.raw_text)

        if match:
            effective_date = match.group(1).strip()
            effective_date = _convert_date(effective_date)
            if effective_date:
                return effective_date.lower()

        return None

    def _extract_planname(self) -> str | None:
        """Extracts the insurance plan name from the provided Document file.

        Assumes the title is located in the first table and it is an only
        paragraph in the cell.

        Returns:
            str: The extracted insurance plan name if found.
            None: If no insurance plan name is found.
        """
        plan_name = ""
        for table in self.document.tables[:1]:
            for cell in table._cells: # pylint: disable=W0212
                for paragraph in cell.paragraphs:
                    if len(table._cells) == len(cell.paragraphs) == 1: # pylint: disable=W0212
                        plan_name = " ".join(paragraph.text.split("\n"))
        if plan_name:
            return plan_name.lower()
        return None

    def _extract_title(self) -> str | None:
        """Extracts the document title from the provided Document file.

        Assumes the title is located in a paragraph styled as "Title".

        Returns:
            str: The extracted title if found.
            None: If no paragraph with the "Title" style is found.
        """

        title_style = None
        title_text = None
        splitted_text = self.raw_text.split("\n")
        for line in splitted_text:
            if line.lower().startswith("title:"):
                title_text = line[len("title:")+1:]
                break

        for paragraph in self.document.paragraphs:
            if "Title" in paragraph.style.name:
                title_style = paragraph.text.lower()
                break


        if title_style:
            return title_style

        return title_text

    def _extract_doc_identifier(self) -> str | None:
        """Extracts the doc_identifier from the filename of the document.

        Searches for a pattern like "KM... or KB...".

        Returns:
            str: The formatted doc_identifier if found.
            None: If the doc_identifier pattern is not found in the filename.
        """
        patterns = [r"km\d+", r"kb\d+"]

        for pattern in patterns:
            match = re.search(pattern, self.filename.lower())
            if not match:
                match = re.search(pattern, self.raw_text.lower())

            if match:
                doc_identifier = match.group().strip()
                return doc_identifier

        return None

    def create_metadata(self) -> dict[str, str]:
        """Method to be implemented by subclasses.

        Generates a dictionary of metadata extracted from the .docx file.

        Returns:
            dict[str, str]: A dictionary containing metadata keys and their
            corresponding values.
        """
        policy_number = self._extract_policy_number()
        effective_date = self._extract_effective_date()
        title = self._extract_title()
        plan_name = self._extract_planname()
        doc_identifier = self._extract_doc_identifier()

        if policy_number:
            filename = f"{policy_number}-{title}"
        else:
            filename = f"{doc_identifier}-{title}"
        filename = re.sub(r"[^\w.-]", "_", filename)
        filename = re.sub(r"__+", "_", filename).rstrip("_")

        metadata = {
            "data_source": "kc",
            "policy_number": policy_number or "",
            "set_number": "",
            "effective_date": effective_date or "",
            "cancellation_date": "",
            "original_filepath": self.filename,
            "section_name": "",
            "plan_name": plan_name or "",
            "policy_title": title or "",
            "url": "",
            "doc_identifier": doc_identifier or "",
            "category_name": "",
            "benefit_id": "",
            "filename": filename or "",
        }
        return metadata


class DefaultDocxChunker:
    """Default chunker class that provides the structure for extracting sections and chunking .docx documents.

    This class defines the core framework for  working with documents, section
    identification, and text chunking.

    Attributes:
        document (docx.Document): The loaded Document object.
        raw_text (str): The raw text content extracted from the .docx document.
        chunk_level (int): The desired level of granularity for chunking (based
          on heading levels, default is 1).
    """

    def __init__(
        self, document: docx.Document, raw_text: str, chunk_level: int = 1
    ):
        self.document = document
        self.raw_text = raw_text
        self.chunk_level = chunk_level

    def extract_sections(self, style: str = "heading") -> list[tuple[str, str]]:
        """Extracts sections from the .docx document based on a specified heading style.

        Args:
            style (str, optional): The style name to use for identifying
              sections. Defaults to "heading".

        Returns:
            list[tuple[str, str]]: A list of tuples where each tuple represents
            (heading_level, section_title) if sections are found.
            None: If no sections matching the style are found.
        """
        sections_names = []
        for paragraph in self.document.paragraphs:
            if style.lower() in paragraph.style.name.lower():
                found_results = re.findall(r"\d+", paragraph.style.name)
                if found_results:
                    level = int(found_results[0])
                    sections_names.append((level, paragraph.text))
        return sections_names

    def get_next_section_index(
        self, sections: list[tuple[str, str]], i: int = 0
    ) -> int:
        """
        Finds the index of the next section within the sections list, considering sections at 
        or below the `chunk_level`.

        Args:
            sections (list[tuple[str, str]]): A list of (heading_level,
              section_title) tuples.
            i (int, optional): The starting index for the search. Defaults to 0.

        Returns:
            int: The index of the next section that is at or below
            `chunk_level`, or -1 if out of list range.
        """

        if i >= len(sections):
            return -1
        while i < len(sections) and sections[i][0] > self.chunk_level:
            i += 1
        return i

    def chunk_the_document(self) -> dict[tuple[int, str], str]:
        """Divides the .docx document into chunks based on the defined sections names.

        Returns:
            dict[tuple[int, str], str]: A dictionary where the keys are tuples
            of (section_id, section_title) and the values are the corresponding
            text chunks.
        """

        sections = self.extract_sections("heading") + [
            (1, "Final"),
        ]
        splitted_text = self.raw_text.split("\n")

        current_section = "Introduction"  # The section placeholder
        current_level = self.chunk_level  # assumed lowest level
        current_text = ""
        # for cases when section names are duplicated
        section_id = 1

        output_file = {}
        i = self.get_next_section_index(sections)

        for line in splitted_text:
            if line.strip() == sections[i][1].strip():
                if current_level == self.chunk_level:
                    output_file[(section_id, current_section)] = current_text
                    section_id += 1
                current_level, current_section = sections[i]
                current_text = ""
                i = self.get_next_section_index(sections, i + 1)
                if i == -1:
                    break
            current_text += f"{line.strip()}\n"

        # save last chunk
        if current_level == self.chunk_level:
            output_file[(section_id, current_section)] = current_text

        return output_file


class CustomKcDocxChunker(DefaultDocxChunker):
    """
    Document chunker class customly created for Custom-KC use case, that provides the structure
    for extracting sections and chunking .docx documents.

    Attributes:
        document (docx.Document): The loaded Document object.
        raw_text (str): The raw text content extracted from the .docx document.
        chunk_level (int): The desired level of granularity for chunking (based
          on heading levels, default is 1).
    """

    def extract_headings_in_table(
        self, style_name: str = "Heading 3"
    ) -> list[str] | None:
        """Extracts sections from the .docx document tables based on a specified heading style.

        Args:
            style_name (str, optional): The style name to use for identifying
              sections. Defaults to "Heading 3".

        Returns:
            list[str]: A list of different sections, if sections are found.
            None: If no sections matching the style are found.
        """
        hdngs = set()
        hldngs_list = []  # use list for order
        for table in self.document.tables:
            for cell in table._cells: # pylint: disable=W0212
                for paragraph in cell.paragraphs:
                    if style_name in paragraph.style.name:
                        if paragraph.text not in hdngs:
                            hldngs_list.append(paragraph.text)
                            hdngs.add(paragraph.text)

        if hldngs_list:
            return hldngs_list
        return None

    def chunk_the_document(self) -> dict[tuple[int, str], str] | None:
        """Divides the .docx document into chunks based on the defined sections names.

        Returns:
            dict[tuple[int, str], str]: A dictionary where the keys are tuples
            of (section_id, section_title) and the values are the corresponding
            text chunks.
        """
        sections = self.extract_sections()
        sections = set([item[1] for item in sections if item[0] == 1])

        items = self.extract_headings_in_table()
        if items is None:
            return None

        splitted_text = self.raw_text.split("\n")
        current_section = items[0]
        section_id = 1
        output_file = {}

        beginning = i = splitted_text.index(current_section)

        for section in items[1:] + ["FinalSection"]:
            current_text = ""
            while i < len(splitted_text):
                if splitted_text[i] == section or splitted_text[i] in sections:
                    ending = i
                    current_text += "\n".join(splitted_text[beginning:ending])
                    output_file[(section_id, current_section)] = current_text
                    section_id += 1
                    current_section = section
                    beginning = i
                    break
                i += 1
        return output_file
